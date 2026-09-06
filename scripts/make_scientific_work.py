"""Сборка научной работы в .docx по правилам конкурса.

    python scripts/make_scientific_work.py

Читает размеченные части из `docs/work/*.md` и собирает
`artifacts/Vertex_scientific_work.docx`: титульный лист по приложению 2,
абстракт на трёх языках, оглавление, исследовательскую часть, заключение,
список источников и дневник.

Оформление: Times New Roman 14 пт, полуторный интервал, поля 2/1/2/2 см,
абзацный отступ 1,25 см, выравнивание по ширине — то, что ожидает
проверяющий, и то, чего не даёт экспорт из markdown.

Разметка строк во входных файлах:
    H1|, H2|      заголовки разделов и подразделов
    P|            абзац (**жирный** внутри поддерживается)
    LIST|         элемент списка
    FORMULA|      формула отдельной строкой, по центру, курсивом
    TABLE_*|      вставка одной из таблиц, собранных ниже из измерений
    FIGURE|       рисунок: путь и подпись через «;»
    TOC|          автообновляемое поле оглавления
    PAGEBREAK|    разрыв страницы
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

PARTS = ["part1_intro.md", "part2_method.md", "part3_conclusion.md"]
SOURCE = Path("docs") / "work"
OUT = Path("artifacts") / "Vertex_scientific_work.docx"

FONT = "Times New Roman"
SIZE = Pt(14)

# --------------------------------------------------------------------------
# Таблицы: числа те же, что в docs/RESULTS.md и artifacts/*.json
# --------------------------------------------------------------------------

TABLES: dict[str, tuple[str, list[list[str]]]] = {
    "TABLE_LADDER": (
        "Таблица 1. Лестница миров: сложность объявлена до проведения прогонов",
        [
            ["Мир", "Какие честные участники добавляются", "Что проверяет ступень"],
            ["W1", "Зарплаты, фриланс", "Проверка прибора: сигнал заведомо есть"],
            ["W2", "Быстро расходующие средства", "Разделяют поведение, по которому узнаётся вывод"],
            ["W3", "Сборы, продавцы, семейные круги", "Негативы той же формы, что и мошенники"],
            ["W4", "Пирамиды и криптовалютные цепочки", "Две схемы на разных временных масштабах"],
            ["W5", "Те же, плюс уклонение организатора", "Цена сокрытия для каждого объекта анализа"],
        ],
    ),
    "TABLE_RESULTS": (
        "Таблица 2. Показатель ROC-AUC по ступеням лестницы, среднее по трём запускам",
        [
            ["Мир", "Счётный уровень", "Групповой уровень"],
            ["W1", "0,955", "0,990"],
            ["W2", "0,941", "0,994"],
            ["W3", "0,923", "0,9996"],
            ["W4", "0,950", "0,9997"],
            ["W5 (уклонение)", "0,966", "оценка отсутствует"],
        ],
    ),
    "TABLE_RARITY": (
        "Таблица 3. Работа детектора при разной доле мошенников, среднее по трём запускам",
        [
            ["Доля мошенников", "Их число в выборке", "ROC-AUC", "Доля настоящих в верхних 10 %", "Проверок на одну находку"],
            ["6,9 %", "950", "0,957", "50 %", "2"],
            ["1,0 %", "126", "0,933", "6,6 %", "15"],
            ["0,5 %", "63", "0,949", "3,3 %", "33"],
            ["0,1 %", "13", "0,894", "0,7 %", "160"],
        ],
    ),
    "TABLE_POLICY": (
        "Таблица 4. Две политики отбора дел при доле мошенников 0,1 %",
        [
            ["Политика", "Сигналов на 1000 счетов", "Доля подтверждений"],
            ["Проверять верхние 10 % списка", "100", "0,008"],
            ["Порог под половину всех мошенников", "0,6", "0,86"],
            ["Порог под четырёх мошенников из пяти", "109", "0,011"],
        ],
    ),
    "TABLE_EVASION": (
        "Таблица 5. Доля обнаруженных групп в зависимости от параметров уклонения",
        [
            ["Источников финансирования", "Банкоматов", "Доля найденных групп", "Медиана перекрытия"],
            ["1", "1", "1,000", "1,000"],
            ["2", "1", "0,956", "0,791"],
            ["3", "1", "0,667", "0,554"],
            ["4", "1", "0,378", "0,417"],
            ["6", "1", "0,100", "0,228"],
            ["1", "2", "0,978", "1,000"],
            ["1", "3", "0,761", "1,000"],
            ["1", "4", "0,761", "1,000"],
            ["6", "4", "0,000", "0,000"],
        ],
    ),
    "TABLE_BASELINE": (
        "Таблица 6. Сравнение с действующими критериями, десять независимых запусков",
        [
            ["Подход", "ROC-AUC", "Разброс"],
            ["Критерии приказа (три из четырёх выразимы)", "0,761", "± 0,008"],
            ["Ансамбль «Vertex» на тех же данных", "0,959", "± 0,009"],
        ],
    ),
    "TABLE_DIARY": (
        "Таблица 7. Дневник выполнения работы",
        [
            ["Этап", "Содержание работы", "Результат"],
            ["1", "Анализ типологий мошенничества рынка Казахстана и действующих критериев", "Формализованы пять структурных подписей"],
            ["2", "Разработка симулятора финансового мира", "Восемь честных популяций, четыре трудных отрицательных"],
            ["3", "Обнаружение утечки ответа в признаки", "Разделены генерация событий и вычисление признаков; ROC-AUC 1,0000 → 0,81"],
            ["4", "Разработка процедуры поиска групп вслепую", "Введено понятие покрытия как потолка полноты"],
            ["5", "Формализация и реализация параметра W через FIFO-сопоставление", "Параметр в интервале [0, 1] с индивидуальным временем удержания"],
            ["6", "Проверка метода одностороннего отбора рёбер (OES)", "Отрицательный результат: сохраняется 0,7206 событий; вызовы удалены"],
            ["7", "Введение purged walk-forward и потолков объектов анализа", "Исключена утечка данных из будущего"],
            ["8", "Построение лестницы из пяти миров", "Сложность объявлена до прогонов, показаны все ступени"],
            ["9", "Обнаружение дефекта первой лестницы", "Мир на 96 % из мошенников; правило закреплено тестом"],
            ["10", "Измерение при доле мошенников 6,9 / 1 / 0,5 / 0,1 %", "Установлена нечувствительность ROC-AUC к редкости"],
            ["11", "Измерение цены уклонения по каждому параметру", "Точка разрушения — три источника финансирования"],
            ["12", "Добавление криптовалютного канала и честных трейдеров", "Четыре типологии из пяти срабатывают в мире"],
            ["13", "Сборка конвейера от событий до очереди дел", "Полный прогон одной командой, 252 автоматических теста"],
        ],
    ),
}


def _style_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = SIZE
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    paragraph_format = normal.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_after = Pt(0)

    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)


def _add_page_numbers(document: Document) -> None:
    """Номер страницы по центру нижнего колонтитула."""
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    for instruction in ("begin", "text", "separate", "end"):
        element = OxmlElement(f"w:fld{'Char' if instruction != 'text' else 'Simple'}")
        if instruction == "text":
            element = OxmlElement("w:instrText")
            element.set(qn("xml:space"), "preserve")
            element.text = " PAGE "
        else:
            element.set(qn("w:fldCharType"), instruction)
        run._r.append(element)
    run.font.name = FONT
    run.font.size = Pt(12)


def _bold_runs(paragraph, text: str) -> None:
    """Разметка **жирным** внутри абзаца."""
    for index, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        run.bold = index % 2 == 1
        run.font.name = FONT
        run.font.size = SIZE


def _paragraph(document: Document, text: str, *, first_line: bool = True):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.space_after = Pt(6)
    _bold_runs(paragraph, text)
    return paragraph


def _heading(document: Document, text: str, level: int):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if level == 2 else WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(14 if level == 2 else 15)
    # Заголовок в стиле, который подхватит автособираемое оглавление.
    paragraph.style = document.styles["Heading 1" if level == 1 else "Heading 2"]
    for child in paragraph.runs:
        child.font.color.rgb = None
        child.font.name = FONT
        child.font.size = Pt(14 if level == 2 else 15)
        child.bold = True
    return paragraph


def _table_of_contents(document: Document) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-2" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Оглавление обновляется в Word: правая кнопка → «Обновить поле»."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, placeholder, end):
        run._r.append(element)


def _table(document: Document, key: str) -> None:
    caption, rows = TABLES[key]
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if column_index or row_index == 0
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(value)
            run.bold = row_index == 0
            run.font.name = FONT
            run.font.size = Pt(12)
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(6)
    note.paragraph_format.space_after = Pt(10)
    run = note.add_run(caption)
    run.italic = True
    run.font.name = FONT
    run.font.size = Pt(12)


def build() -> Path:
    document = Document()
    _style_document(document)
    _add_page_numbers(document)

    for name in PARTS:
        for line in (SOURCE / name).read_text(encoding="utf-8").splitlines():
            if not line.strip():
                continue
            tag, _, text = line.partition("|")
            if tag == "TITLE":
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(text)
                run.bold = True
                run.font.name = FONT
                run.font.size = Pt(14)
            elif tag == "TITLE2":
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(text)
                run.bold = True
                run.font.name = FONT
                run.font.size = Pt(16)
            elif tag == "CENTER":
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(text)
                run.font.name = FONT
                run.font.size = Pt(14)
            elif tag == "GAP":
                document.add_paragraph()
            elif tag == "PAGEBREAK":
                document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            elif tag == "H1":
                _heading(document, text, 1)
            elif tag == "H2":
                _heading(document, text, 2)
            elif tag == "P":
                _paragraph(document, text)
            elif tag == "LIST":
                paragraph = _paragraph(document, text, first_line=False)
                paragraph.paragraph_format.left_indent = Cm(1.0)
            elif tag == "REF":
                paragraph = _paragraph(document, text, first_line=False)
                paragraph.paragraph_format.left_indent = Cm(0.75)
                paragraph.paragraph_format.space_after = Pt(4)
            elif tag == "FORMULA":
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(6)
                paragraph.paragraph_format.space_after = Pt(6)
                run = paragraph.add_run(text)
                run.italic = True
                run.font.name = FONT
                run.font.size = SIZE
            elif tag == "FIGURE":
                image, _, caption = text.partition(";")
                document.add_picture(image.strip(), width=Cm(16))
                document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                signature = document.add_paragraph()
                signature.alignment = WD_ALIGN_PARAGRAPH.CENTER
                signature.paragraph_format.space_before = Pt(4)
                signature.paragraph_format.space_after = Pt(10)
                run = signature.add_run(caption.strip())
                run.italic = True
                run.font.name = FONT
                run.font.size = Pt(12)
            elif tag == "TOC":
                _table_of_contents(document)
            elif tag in TABLES:
                _table(document, tag)
            else:
                raise ValueError(f"неизвестная разметка: {tag}")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Готово: {path}  ({path.stat().st_size / 1024:.0f} КБ)")
